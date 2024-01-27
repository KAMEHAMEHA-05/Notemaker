# -*- coding: utf-8 -*-
"""
Created on Sun Nov 21 20:11:02 2021

@author: Admin
"""

from pygame import display
from pygame import init
from pygame import image
from pygame import Rect
from pygame import draw
from pygame import font
from pygame import event
from pygame import MOUSEBUTTONDOWN
from pygame import quit
import os
from docx import Document
from pyautogui import alert
import math
import tkinter as tk
from tkinter import filedialog
from nltk import sent_tokenize, word_tokenize, PorterStemmer
from nltk.corpus import stopwords

def _create_frequency_matrix(sentences):
    frequency_matrix = {}
    stopWords = set(stopwords.words("english"))
    ps = PorterStemmer()

    for sent in sentences:
        freq_table = {}
        words = word_tokenize(sent)
        for word in words:
            word = word.lower()
            word = ps.stem(word)
            if word in stopWords:
                continue

            if word in freq_table:
                freq_table[word] += 1
            else:
                freq_table[word] = 1

        frequency_matrix[sent[:15]] = freq_table

    return frequency_matrix

def _create_tf_matrix(freq_matrix):
    tf_matrix = {}

    for sent, f_table in freq_matrix.items():
        tf_table = {}

        count_words_in_sentence = len(f_table)
        for word, count in f_table.items():
            tf_table[word] = count / count_words_in_sentence

        tf_matrix[sent] = tf_table

    return tf_matrix

def _create_documents_per_words(freq_matrix):
    word_per_doc_table = {}

    for sent, f_table in freq_matrix.items():
        for word, count in f_table.items():
            if word in word_per_doc_table:
                word_per_doc_table[word] += 1
            else:
                word_per_doc_table[word] = 1

    return word_per_doc_table

def _create_idf_matrix(freq_matrix, count_doc_per_words, total_documents):
    idf_matrix = {}

    for sent, f_table in freq_matrix.items():
        idf_table = {}

        for word in f_table.keys():
            idf_table[word] = math.log10(total_documents / float(count_doc_per_words[word]))

        idf_matrix[sent] = idf_table

    return idf_matrix

def _create_tf_idf_matrix(tf_matrix, idf_matrix):
    tf_idf_matrix = {}

    for (sent1, f_table1), (sent2, f_table2) in zip(tf_matrix.items(), idf_matrix.items()):

        tf_idf_table = {}

        for (word1, value1), (word2, value2) in zip(f_table1.items(),
                                                    f_table2.items()):  # here, keys are the same in both the table
            tf_idf_table[word1] = float(value1 * value2)

        tf_idf_matrix[sent1] = tf_idf_table

    return tf_idf_matrix

def _score_sentences(tf_idf_matrix) -> dict:

    sentenceValue = {}

    for sent, f_table in tf_idf_matrix.items():
        total_score_per_sentence = 0

        count_words_in_sentence = len(f_table)
        for word, score in f_table.items():
            total_score_per_sentence += score

        sentenceValue[sent] = total_score_per_sentence / count_words_in_sentence

    return sentenceValue

def _find_average_score(sentenceValue) -> int:

    sumValues = 0
    for entry in sentenceValue:
        sumValues += sentenceValue[entry]

    # Average value of a sentence from original summary_text
    average = (sumValues / len(sentenceValue))

    return average

def _generate_summary(sentences, sentenceValue, threshold):
    sentence_count = 0
    summary = ''

    for sentence in sentences:
        if sentence[:15] in sentenceValue and sentenceValue[sentence[:15]] >= (threshold):
            summary += " " + sentence
            sentence_count += 1

    return summary

curr_dir = os.getcwd()
curr_dir = curr_dir[:-8]
notes = os.path.join(curr_dir, "Notes")
os.chdir(notes)

class Notemaker():
    def __init__(self, gameWindow):
        init()

        self.g_w = gameWindow

    def main(self):
        self.game_setup()

        running = True
        while running:
             for ev in event.get():
                 if ev.type == quit():
                     running = False

    def game_setup(self):
        init()
        self.g_w = display.set_mode((275, 385))
        icon = image.load(r"C:\Users\Admin\Downloads\icon.ico")
        display.set_caption("New Note")
        display.set_icon(icon)
        req = "Requirements\e_new_note_bg.png"
        n_bg_a = os.path.join(curr_dir,req)
        notes_bg_a = image.load(n_bg_a)
        self.g_w.blit(notes_bg_a, (0,0))
        browse = Rect(9, 67, 78, 23)
        location = Rect(9, 153, 87, 23)
        browse_box = Rect(12, 97, 249, 23)
        location_box = Rect(12, 184, 249, 24)
        next_button = Rect(86, 295, 102, 40)
        draw.rect(self.g_w, (252, 252, 252), browse, 1)
        draw.rect(self.g_w, (252, 252, 252), location, 1)
        draw.rect(self.g_w, (252, 252, 252), next_button, 1)
        display.flip()
        
        pressed_b = 0
        pressed_l = 0
        pressed_n = 0
        
        running = True
        while running : 
            for eve in event.get():
                number = eve.type
                if eve.type == (MOUSEBUTTONDOWN):
                     if browse.collidepoint(eve.pos) and pressed_b == 0 :
                         root = tk.Tk()
                         file_path = filedialog.askopenfilename()
                         root.withdraw()
                         file_name = os.path.basename(file_path)
                         file = open(file_path,"r")
                         Font_ = font.Font(font.get_default_font(), 13)
                         text_surface = Font_.render(file_name, True, (0, 0, 0))
                         self.g_w.blit(text_surface, (13,99))
                         display.flip()
                         pressed_b =  0
                         
                     if location.collidepoint(eve.pos) and pressed_l == 0 :
                         location_ = filedialog.askdirectory()
                         root.withdraw()
                         directory = os.path.basename(location_)
                         Font_ = font.Font(font.get_default_font(), 13)
                         text_surface = Font_.render(directory, True, (0, 0, 0))
                         self.g_w.blit(text_surface, (13,187))
                         display.flip()
                         pressed_l = 0
                         
                     if next_button.collidepoint(eve.pos) and pressed_n == 0 :
                        # Seperate out sentences
                        sentences = sent_tokenize(file.read())
                        original_text = " ".join(sentences)
                        total_documents = len(sentences)
                        
                        # Create Frequency matrix of the words in each sentence.
                        freq_matrix = _create_frequency_matrix(sentences)
                        
                        # Calculate TF and create matrix
                        tf_matrix = _create_tf_matrix(freq_matrix)
                        
                        # create table for docs/words
                        count_doc_per_words = _create_documents_per_words(freq_matrix)
                        
                        # Calculate IDF and create matrix
                        idf_matrix = _create_idf_matrix(freq_matrix, count_doc_per_words, total_documents)
                        
                        # Calculate TF-IDF and create matrix
                        tf_idf_matrix = _create_tf_idf_matrix(tf_matrix, idf_matrix)
                        
                        # score sentences
                        sentence_scores = _score_sentences(tf_idf_matrix)
                        
                        # Determine threshold
                        threshold = _find_average_score(sentence_scores)
                     
                        # Create notes
                        summary = _generate_summary(sentences, sentence_scores, 0.845 * threshold)
                        summary_words_ = "No of words in summary :", str(len(summary))
                        summary_words = " ".join(summary_words_)
                        original_words_ = "No of words in original text :", str(len(original_text))
                        original_words = " ".join(original_words_)
                        #print(summary_words)
                        #print(original_words)
                        
                        name = file_name[:-4]
                        
                        document = Document()
                        document.add_heading("Notes of "+name, 0)
                        document.add_paragraph(summary)
                        document.add_paragraph("====================================================")
                        document.add_paragraph(summary_words)
                        document.add_paragraph(original_words)
                        document.save("notes_"+name+".docx")
                        word_file_ = "notes_"+name+".docx"
                        word_file = os.path.join(location_,word_file_)
                        document.save(word_file)
                        alert(".docx file of the notes has been saved in the location")
                        draw.rect(self.g_w, (220, 220, 220), browse_box)
                        draw.rect(self.g_w, (220, 220, 220), location_box)
                        display.flip()
                        pressed_n = 0
                 
                if number == 256 :
                    running = False
                
        quit()
                
game = Notemaker(None)
game.main()
